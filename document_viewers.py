import streamlit as st
from typing import Optional, Dict, Any
import docx
import io

def get_theme_colors(theme: str) -> Dict[str, str]:
    """Get color scheme for different themes"""
    themes = {
        "Light": {
            "bg": "#FFFFFF",
            "text": "#000000",
            "heading": "#333333",
            "border": "#E0E0E0"
        },
        "Dark": {
            "bg": "#1E1E1E",
            "text": "#E0E0E0",
            "heading": "#FFFFFF",
            "border": "#333333"
        },
        "Sepia": {
            "bg": "#F4ECD8",
            "text": "#5C4B51",
            "heading": "#2C1810",
            "border": "#C4A484"
        }
    }
    return themes.get(theme, themes["Light"])

def pdf_viewer(
    binary_content: bytes,
    width: int = 800,
    enable_text: bool = True,
    pages_vertical_spacing: int = 2,
    resolution_boost: int = 1,
    annotation_outline_size: int = 1
) -> None:
    """
    PDF viewer with enhanced controls and consistent styling.
    
    Args:
        binary_content: Raw bytes of the PDF file
        width: Display width in pixels
        enable_text: Whether to enable text selection
        pages_vertical_spacing: Space between pages in pixels
        resolution_boost: Quality multiplier for rendering
        annotation_outline_size: Thickness of annotation borders in pixels
    """
    try:
        from streamlit_pdf_viewer import pdf_viewer
        pdf_viewer(
            input=binary_content,
            width=width,
            pages_vertical_spacing=pages_vertical_spacing,
            annotation_outline_size=annotation_outline_size,
            render_text=enable_text,
            resolution_boost=resolution_boost
        )
    except Exception as e:
        st.error(f"Error displaying PDF: {str(e)}")
        print(f"PDF viewer error details: {str(e)}")

def docx_viewer(
    binary_content: bytes,
    width: int = 800,
    theme: str = "Light",
    font_size: int = 16,
    show_headings: bool = True,
    show_paragraphs: bool = True,
    show_lists: bool = True
) -> None:
    """
    Dedicated viewer for DOCX files with rich formatting support.
    
    Args:
        binary_content: Raw bytes of the DOCX file
        width: Display width in pixels
        theme: Color theme ("Light", "Dark", "Sepia")
        font_size: Base font size in pixels
        show_headings: Whether to show heading structure
        show_paragraphs: Whether to show paragraph structure
        show_lists: Whether to show list structure
    """
    try:
        # Create document object
        doc = docx.Document(io.BytesIO(binary_content))
        theme_colors = get_theme_colors(theme)
        
        # Create container for document
        doc_container = st.container()
        
        with doc_container:
            # Main document container
            st.markdown(f"""
                <div style="
                    background-color: {theme_colors['bg']};
                    padding: 2rem;
                    border-radius: 8px;
                    width: {width}px;
                    margin: 0 auto;
                    color: {theme_colors['text']};
                    font-family: 'Georgia', serif;
                    border: 1px solid {theme_colors['border']};
                    box-shadow: 0 2px 4px rgba(0,0,0,0.1);
                ">
            """, unsafe_allow_html=True)
            
            # Process each paragraph
            for para in doc.paragraphs:
                if not para.text.strip():
                    continue
                
                style_name = para.style.name.lower()
                
                # Handle headings
                if 'heading' in style_name and show_headings:
                    level = style_name[-1] if style_name[-1].isdigit() else '2'
                    font_size_heading = font_size + (4 - int(level)) * 2
                    
                    st.markdown(f"""
                        <h{level} style="
                            color: {theme_colors['heading']};
                            font-size: {font_size_heading}px;
                            margin: 1em 0 0.5em 0;
                            font-weight: bold;
                            border-bottom: 1px solid {theme_colors['border']};
                            padding-bottom: 0.2em;
                        ">
                            {para.text}
                        </h{level}>
                    """, unsafe_allow_html=True)
                
                # Handle lists
                elif show_lists and (para.text.strip().startswith('-') or 
                                   para.text.strip().startswith('•') or
                                   para.text.strip().startswith('*')):
                    st.markdown(f"""
                        <ul style="
                            margin-left: 1.5em;
                            margin-bottom: 0.5em;
                            font-size: {font_size}px;
                            color: {theme_colors['text']};
                        ">
                            <li style="margin-bottom: 0.3em;">
                                {para.text.lstrip('- •*')}
                            </li>
                        </ul>
                    """, unsafe_allow_html=True)
                
                # Handle regular paragraphs
                elif show_paragraphs:
                    st.markdown(f"""
                        <p style="
                            margin-bottom: 1.2em;
                            font-size: {font_size}px;
                            line-height: 1.6;
                            color: {theme_colors['text']};
                            text-align: justify;
                        ">
                            {para.text}
                        </p>
                    """, unsafe_allow_html=True)
            
            # Close main container
            st.markdown("</div>", unsafe_allow_html=True)
            
    except Exception as e:
        st.error(f"Error displaying DOCX: {str(e)}")
        print(f"DOCX viewer error details: {str(e)}")


def text_viewer(
    binary_content: bytes,
    width: int = 800,
    theme: str = "Light",
    font_size: int = 16,
    line_spacing: float = 1.6
) -> None:
    """
    Dedicated viewer for text files with enhanced readability.
    
    Args:
        binary_content: Raw bytes of the text file
        width: Display width in pixels
        theme: Color theme ("Light", "Dark", "Sepia")
        font_size: Font size in pixels
        line_spacing: Line height multiplier
    """
    try:
        content = binary_content.decode('utf-8')
        theme_colors = get_theme_colors(theme)
        
        # Split content into paragraphs
        paragraphs = [p for p in content.split('\n\n') if p.strip()]
        
        # Container for text content
        text_container = st.container()
        
        with text_container:
            st.markdown(f"""
                <div style="
                    background-color: {theme_colors['bg']};
                    padding: 2rem;
                    border-radius: 8px;
                    width: {width}px;
                    margin: 0 auto;
                    color: {theme_colors['text']};
                    font-family: 'Georgia', serif;
                    border: 1px solid {theme_colors['border']};
                    box-shadow: 0 2px 4px rgba(0,0,0,0.1);
                ">
            """, unsafe_allow_html=True)
            
            for para in paragraphs:
                # Check if paragraph is a header (starts with #)
                if para.strip().startswith('#'):
                    level = len(para.split()[0])  # Count number of # symbols
                    text = ' '.join(para.split()[1:])
                    font_size_heading = font_size + (4 - level) * 2
                    
                    st.markdown(f"""
                        <h{level} style="
                            color: {theme_colors['heading']};
                            font-size: {font_size_heading}px;
                            margin: 1em 0 0.5em 0;
                            font-weight: bold;
                            border-bottom: 1px solid {theme_colors['border']};
                            padding-bottom: 0.2em;
                        ">{text}</h{level}>
                    """, unsafe_allow_html=True)
                    
                # Check if paragraph is a list item
                elif para.strip().startswith(('-', '*', '•')):
                    st.markdown(f"""
                        <ul style="
                            margin-left: 1.5em;
                            margin-bottom: 0.5em;
                            font-size: {font_size}px;
                            color: {theme_colors['text']};
                        ">
                            <li style="margin-bottom: 0.3em;">
                                {para.strip().lstrip('- *•')}
                            </li>
                        </ul>
                    """, unsafe_allow_html=True)
                    
                # Regular paragraph
                else:
                    st.markdown(f"""
                        <p style="
                            margin-bottom: 1.2em;
                            font-size: {font_size}px;
                            line-height: {line_spacing};
                            color: {theme_colors['text']};
                            text-align: justify;
                        ">{para}</p>
                    """, unsafe_allow_html=True)
            
            st.markdown("</div>", unsafe_allow_html=True)
            
    except Exception as e:
        st.error(f"Error displaying text: {str(e)}")

def get_theme_colors(theme: str) -> Dict[str, str]:
    """Get color scheme for different themes"""
    themes = {
        "Light": {
            "bg": "#FFFFFF",
            "text": "#000000",
            "heading": "#333333",
            "border": "#E0E0E0"
        },
        "Dark": {
            "bg": "#1E1E1E",
            "text": "#E0E0E0",
            "heading": "#FFFFFF",
            "border": "#333333"
        },
        "Sepia": {
            "bg": "#F4ECD8",
            "text": "#5C4B51",
            "heading": "#2C1810",
            "border": "#C4A484"
        }
    }
    return themes.get(theme, themes["Light"])

def docx_viewer(
    binary_content: bytes,
    width: int = 800,
    theme: str = "Light",
    font_size: int = 16,
    show_headings: bool = True,
    show_paragraphs: bool = True,
    show_lists: bool = True
) -> None:
    """
    Dedicated viewer for DOCX files with rich formatting support.
    
    Args:
        binary_content: Raw bytes of the DOCX file
        width: Display width in pixels
        theme: Color theme ("Light", "Dark", "Sepia")
        font_size: Base font size in pixels
        show_headings: Whether to show heading structure
        show_paragraphs: Whether to show paragraph structure
        show_lists: Whether to show list structure
    """
    try:
        doc = docx.Document(io.BytesIO(binary_content))
        theme_colors = get_theme_colors(theme)
        
        # Container for document content
        doc_container = st.container()
        
        with doc_container:
            st.markdown(f"""
                <div style="
                    background-color: {theme_colors['bg']};
                    padding: 2rem;
                    border-radius: 8px;
                    width: {width}px;
                    margin: 0 auto;
                    color: {theme_colors['text']};
                    font-family: 'Georgia', serif;
                    border: 1px solid {theme_colors['border']};
                    box-shadow: 0 2px 4px rgba(0,0,0,0.1);
                ">
            """, unsafe_allow_html=True)
            
            for para in doc.paragraphs:
                if not para.text.strip():
                    continue
                
                style_name = para.style.name.lower()
                
                # Handle headings
                if 'heading' in style_name and show_headings:
                    level = style_name[-1] if style_name[-1].isdigit() else '2'
                    font_size_heading = font_size + (4 - int(level)) * 2
                    
                    st.markdown(f"""
                        <h{level} style="
                            color: {theme_colors['heading']};
                            font-size: {font_size_heading}px;
                            margin: 1em 0 0.5em 0;
                            font-weight: bold;
                            border-bottom: 1px solid {theme_colors['border']};
                            padding-bottom: 0.2em;
                        ">{para.text}</h{level}>
                    """, unsafe_allow_html=True)
                
                # Handle lists
                elif show_lists and (para.text.strip().startswith('-') or 
                                   para.text.strip().startswith('•') or
                                   para.text.strip().startswith('*')):
                    st.markdown(f"""
                        <ul style="
                            margin-left: 1.5em;
                            margin-bottom: 0.5em;
                            font-size: {font_size}px;
                            color: {theme_colors['text']};
                        ">
                            <li style="margin-bottom: 0.3em;">{para.text.lstrip('- •*')}</li>
                        </ul>
                    """, unsafe_allow_html=True)
                
                # Handle regular paragraphs
                elif show_paragraphs:
                    st.markdown(f"""
                        <p style="
                            margin-bottom: 1.2em;
                            font-size: {font_size}px;
                            line-height: 1.6;
                            color: {theme_colors['text']};
                            text-align: justify;
                        ">{para.text}</p>
                    """, unsafe_allow_html=True)
            
            st.markdown("</div>", unsafe_allow_html=True)
            
    except Exception as e:
        st.error(f"Error displaying DOCX: {str(e)}")

def text_viewer(
    binary_content: bytes,
    width: int = 800,
    theme: str = "Light",
    font_size: int = 16,
    line_spacing: float = 1.6
) -> None:
    """
    Dedicated viewer for text files with enhanced readability.
    
    Args:
        binary_content: Raw bytes of the text file
        width: Display width in pixels
        theme: Color theme ("Light", "Dark", "Sepia")
        font_size: Font size in pixels
        line_spacing: Line height multiplier
    """
    try:
        content = binary_content.decode('utf-8')
        theme_colors = get_theme_colors(theme)
        
        # Split content into paragraphs
        paragraphs = [p for p in content.split('\n\n') if p.strip()]
        
        # Container for text content
        text_container = st.container()
        
        with text_container:
            st.markdown(f"""
                <div style="
                    background-color: {theme_colors['bg']};
                    padding: 2rem;
                    border-radius: 8px;
                    width: {width}px;
                    margin: 0 auto;
                    color: {theme_colors['text']};
                    font-family: 'Georgia', serif;
                    border: 1px solid {theme_colors['border']};
                    box-shadow: 0 2px 4px rgba(0,0,0,0.1);
                ">
            """, unsafe_allow_html=True)
            
            for para in paragraphs:
                # Check if paragraph is a header (starts with #)
                if para.strip().startswith('#'):
                    level = len(para.split()[0])  # Count number of # symbols
                    text = ' '.join(para.split()[1:])
                    font_size_heading = font_size + (4 - level) * 2
                    
                    st.markdown(f"""
                        <h{level} style="
                            color: {theme_colors['heading']};
                            font-size: {font_size_heading}px;
                            margin: 1em 0 0.5em 0;
                            font-weight: bold;
                            border-bottom: 1px solid {theme_colors['border']};
                            padding-bottom: 0.2em;
                        ">{text}</h{level}>
                    """, unsafe_allow_html=True)
                    
                # Check if paragraph is a list item
                elif para.strip().startswith(('-', '*', '•')):
                    st.markdown(f"""
                        <ul style="
                            margin-left: 1.5em;
                            margin-bottom: 0.5em;
                            font-size: {font_size}px;
                            color: {theme_colors['text']};
                        ">
                            <li style="margin-bottom: 0.3em;">
                                {para.strip().lstrip('- *•')}
                            </li>
                        </ul>
                    """, unsafe_allow_html=True)
                    
                # Regular paragraph
                else:
                    st.markdown(f"""
                        <p style="
                            margin-bottom: 1.2em;
                            font-size: {font_size}px;
                            line-height: {line_spacing};
                            color: {theme_colors['text']};
                            text-align: justify;
                        ">{para}</p>
                    """, unsafe_allow_html=True)
            
            st.markdown("</div>", unsafe_allow_html=True)
            
    except Exception as e:
        st.error(f"Error displaying text: {str(e)}")
